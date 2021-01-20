from docx import Document
from docx.shared import Inches

document = Document()

# Profile Picture
document.add_picture('me.png', width=Inches(2.0))

# Name, Phone & Email details
name = input('What is your name? ')
phone_number = input('What is your Phone Number? ')
email = input('What is your email? ')

document.add_paragraph(
    name.title() + ' | ' + phone_number + ' | ' + email.lower())

#About Me
document.add_heading('About Me')
document.add_paragraph(input('Tell about Yourself? ').capitalize())

#Education
document.add_heading('Education')
p = document.add_paragraph()

university = input('Enter University? ')
from_date = input('From Date? ')
to_date = input('To Date? ')
degree = input('Enter Degree Name? ')
major = input('Enter Major of Degree? ')
minor = input('Enter Minor of Degree? ')

p.add_run(university.title() + '\n').bold = True
p.add_run(from_date + ' - ' + to_date + '\n').italic = True
p.add_run(degree.title() + '\n')
p.add_run(major.capitalize() + ' & ' + minor.capitalize())

#More Universities
while True:
    has_more_universities = input('Do you have more Universities to mention? Yes or No? ')
    if has_more_universities.lower() == 'yes':

        p = document.add_paragraph()

        university = input('Enter University? ')
        from_date = input('From Date? ')
        to_date = input('To Date? ')
        degree = input('Enter Degree Name? ')
        major = input('Enter Major of Degree? ')
        minor = input('Enter Minor of Degree? ')

        p.add_run(university.title() + '\n').bold = True
        p.add_run(from_date + ' - ' + to_date + '\n').italic = True
        p.add_run(degree.title() + '\n')
        p.add_run(major.capitalize() + ' & ' + minor.capitalize())

    else:
        break

#Work Experience
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter Company? ')
from_date = input('From Date? ')
to_date = input('To Date? ')
p.add_run(company.title() + '\n').bold = True
p.add_run(from_date + ' - ' + to_date + '\n').italic = True

experience_details = input('Describe your experience at ' + company + '? ')
p.add_run(experience_details.capitalize())

#More Experiences
while True:
    has_more_experiences = input('Do you have more experiences? Yes or No? ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter Company? ')
        from_date = input('From Date? ')
        to_date = input('To Date? ')
        p.add_run(company.title() + '\n').bold = True
        p.add_run(from_date + ' - ' + to_date + '\n').italic = True

        experience_details = input('Describe your experience at ' + company + '? ')
        p.add_run(experience_details.capitalize())

    else:
        break

#Skills
p = document.add_heading('Skills')
p = document.add_paragraph()

skills = input('Enter Skills? ')
p.add_run(skills.title()).bold = True
p.style = 'List Bullet'

#More Skills
while True:
    has_more_skills = input('Do you have more Skills to mention? Yes or No? ')
    if has_more_skills.lower() == 'yes':

        p = document.add_paragraph()

        skills = input('Enter Skills? ')
        p.add_run(skills.title()).bold = True
        p.style = 'List Bullet'

    else:
        break
    
#Footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using Shahzeb’s Code in Python =)"

document.save('cv.docx')